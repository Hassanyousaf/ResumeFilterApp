import os
import re
import pdfplumber
import pytesseract
from docx import Document
from pdf2image import convert_from_path
from collections import defaultdict
from datetime import datetime

class ResumeFilter:
    def __init__(self, job_description, mandatory_keywords, optional_keywords, min_experience):
        """
        Initialize with:
        - job_description: String of job requirements
        - mandatory_keywords: List of must-have terms
        - optional_keywords: List of nice-to-have terms
        - min_experience: Minimum years of experience required (float)
        """
        self.job_description = job_description
        self.mandatory = [kw.lower() for kw in mandatory_keywords]
        self.optional = [kw.lower() for kw in optional_keywords]
        self.min_experience = min_experience
        
    def extract_text(self, filepath):
        """Extract text from PDF or DOCX file (whole document)"""
        text = ""
        try:
            if filepath.lower().endswith('.pdf'):
                # First try pdfplumber
                with pdfplumber.open(filepath) as pdf:
                    text = "\n".join([page.extract_text() or "" for page in pdf.pages])
                
                # Fallback to OCR if needed
                if not text.strip():
                    images = convert_from_path(filepath)
                    text = "\n".join([pytesseract.image_to_string(img) for img in images])
            
            elif filepath.lower().endswith(('.docx', '.doc')):
                doc = Document(filepath)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text.append(cell.text)
                text = "\n".join(full_text)
                
        except Exception as e:
            print(f"Error processing {filepath}: {e}")
        
        return text.lower()  # Normalize to lowercase

    def extract_experience(self, text):
        """
        Extract total years of experience from resume text
        Returns: float (years of experience) or None if not found
        """
        # Common patterns for experience
        patterns = [
            r'(\d+\.?\d*)\s*(years?|yrs?)\s*(experience|exp)',
            r'experience\s*:\s*(\d+\.?\d*)\s*(years?|yrs?)',
            r'(\d+)\+?\s*(years?|yrs?)\s*in\s*.*(experience|exp)',
            r'(\d+\.?\d*)\s*(years?|yrs?)\s*professional',
            r'(\d+)\s*(years?|yrs?)\s*relevant'
        ]
        
        # Search for explicit experience statements
        for pattern in patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                try:
                    return float(match.group(1))
                except:
                    continue
        
        # If no explicit experience found, try to calculate from dates
        date_pattern = r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s*\d{4}\s*[-–—]\s*(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s*\d{4}|\bpresent\b'
        date_ranges = re.findall(date_pattern, text, flags=re.IGNORECASE)
        
        total_experience = 0
        valid_ranges = 0
        
        for date_range in date_ranges:
            try:
                # Split date range
                dates = re.split(r'\s*[-–—]\s*', date_range)
                if len(dates) != 2:
                    continue
                    
                # Parse start and end dates
                start_date = dates[0].strip()
                end_date = dates[1].strip()
                
                # Handle "Present" end date
                if end_date.lower() == "present":
                    end_date = datetime.now().strftime("%b %Y")
                
                # Parse dates
                start = datetime.strptime(start_date[:3] + " " + start_date[-4:], "%b %Y")
                end = datetime.strptime(end_date[:3] + " " + end_date[-4:], "%b %Y")
                
                # Calculate duration in years
                duration = (end - start).days / 365.25
                if duration > 0:
                    total_experience += duration
                    valid_ranges += 1
            except Exception as e:
                continue
        
        # Return average if we found multiple ranges
        if valid_ranges > 0:
            return total_experience
        
        return None

    def analyze_resume(self, text):
        """
        Analyze entire resume text and return:
        - missing_mandatory: List of missing mandatory keywords
        - keyword_counts: Dictionary of keyword frequencies
        - experience: Extracted years of experience
        - score: Total weighted score
        """
        # Check for mandatory keywords
        missing_mandatory = []
        keyword_counts = defaultdict(int)
        found_sections = defaultdict(list)
        
        # Search for mandatory keywords throughout document
        for kw in self.mandatory:
            pattern = r'\b' + re.escape(kw) + r'\b'
            matches = re.finditer(pattern, text)
            count = 0
            
            for match in matches:
                count += 1
                # Capture surrounding context
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 50)
                context = text[start:end].replace('\n', ' ').strip()
                found_sections[kw].append(context)
            
            if count == 0:
                missing_mandatory.append(kw)
            keyword_counts[kw] = count
        
        # Search for optional keywords
        for kw in self.optional:
            pattern = r'\b' + re.escape(kw) + r'\b'
            keyword_counts[kw] = len(re.findall(pattern, text))
        
        # Extract experience
        experience = self.extract_experience(text)
        
        # Calculate score (mandatory 3x, optional 1x)
        mandatory_score = sum(keyword_counts[kw] * 3 for kw in self.mandatory)
        optional_score = sum(keyword_counts[kw] for kw in self.optional)
        
        # Experience bonus (up to 20% of total possible score)
        exp_bonus = 0
        if experience is not None and experience >= self.min_experience:
            exp_bonus = min(experience / self.min_experience, 1.2) * (mandatory_score + optional_score) * 0.2
        
        total_score = mandatory_score + optional_score + exp_bonus
        
        return {
            'missing_mandatory': missing_mandatory,
            'keyword_counts': dict(keyword_counts),
            'found_sections': dict(found_sections),
            'experience': experience,
            'score': total_score
        }

    def process_resumes(self, resume_dir):
        """Process all resumes in directory and return ranked results"""
        results = []
        
        for filename in os.listdir(resume_dir):
            if not filename.lower().endswith(('.pdf', '.docx', '.doc')):
                continue
                
            filepath = os.path.join(resume_dir, filename)
            text = self.extract_text(filepath)
            
            if not text.strip():
                continue  # Skip empty files
                
            analysis = self.analyze_resume(text)
            
            # Only include resumes with all mandatory keywords
            if not analysis['missing_mandatory']:
                # Check experience requirement
                exp_met = analysis['experience'] is not None and analysis['experience'] >= self.min_experience
                
                results.append({
                    'filename': filename,
                    'score': analysis['score'],
                    'experience': analysis['experience'],
                    'experience_met': exp_met,
                    'missing_mandatory': analysis['missing_mandatory'],
                    'keyword_counts': analysis['keyword_counts'],
                    'found_sections': analysis['found_sections']
                })
        
        # Sort by score (highest first), then by experience met
        results.sort(key=lambda x: (-x['score'], -x['experience_met']))
        return results

# Usage
if __name__ == "__main__":
    job_description = """
    Senior Python Developer Position:
    - Minimum 5 years experience required
    - Must have: Python, Machine Learning, NLP
    - Nice to have: Django, Cloud, TensorFlow
    """
    
    # 2. Define requirements
    mandatory_keywords = ['python', 'machine learning', 'nlp']
    optional_keywords = ['django', 'flask', 'cloud', 'aws', 'tensorflow', 'pytorch']
    min_experience = 5.0  
    
    # 3. Initialize filter
    resume_filter = ResumeFilter(job_description, mandatory_keywords, optional_keywords, min_experience)
    
    # 4. Process resumes
    resume_directory = "resumes"  
    results = resume_filter.process_resumes(resume_directory)
    
    # 5. Display results
    print(f"\n{'='*50}\nJob Requirements:\n{job_description}\n{'='*50}")
    print(f"\nProcessing resumes from: {resume_directory}")
    print(f"Minimum Experience Required: {min_experience} years\n")
    
    if not results:
        print("No resumes matched all requirements")
    else:
        print(f"Found {len(results)} qualified resumes:\n")
        for rank, resume in enumerate(results, 1):
            exp_status = "✓" if resume['experience_met'] else f"✗ (Found: {resume['experience'] or 'N/A'} yrs)"
            
            print(f"{rank}. {resume['filename']} (Score: {resume['score']:.1f})")
            print(f"   Experience: {exp_status}")
            print(f"   Missing Mandatory: {', '.join(resume['missing_mandatory']) or 'None'}")
            
            # Print top keywords with context
            print("   Keyword Evidence:")
            for kw in mandatory_keywords:
                count = resume['keyword_counts'].get(kw, 0)
                if count > 0:
                    print(f"     - {kw} ({count} occurrences):")
                    for i, context in enumerate(resume['found_sections'].get(kw, [])[:2], 1):
                        print(f"       {i}. {context[:100]}...")
            print("\n" + "-"*50 + "\n")


