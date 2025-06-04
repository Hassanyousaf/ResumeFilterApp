import os
import re
import pdfplumber
import pytesseract
from docx import Document
from pdf2image import convert_from_path
from collections import defaultdict
from datetime import datetime
import gradio as gr
import tempfile
import shutil

class ResumeFilter:
    def __init__(self, job_description, mandatory_keywords, optional_keywords, min_experience):
        """
        Initialize with:
        - job_description: String of job requirements
        - mandatory_keywords: List of must-have terms
        - optional_keywords: List of nice-to-have terms
        - min_experience: Minimum years of experience required (float)
        """
        self.job_description = job_description
        self.mandatory = [kw.lower() for kw in mandatory_keywords]
        self.optional = [kw.lower() for kw in optional_keywords]
        self.min_experience = min_experience
        
    def extract_text(self, filepath):
        """Extract text from PDF or DOCX file (whole document)"""
        text = ""
        try:
            if filepath.lower().endswith('.pdf'):
                # First try pdfplumber
                with pdfplumber.open(filepath) as pdf:
                    text = "\n".join([page.extract_text() or "" for page in pdf.pages])
                
                # Fallback to OCR if needed
                if not text.strip():
                    images = convert_from_path(filepath)
                    text = "\n".join([pytesseract.image_to_string(img) for img in images])
            
            elif filepath.lower().endswith(('.docx', '.doc')):
                doc = Document(filepath)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text.append(cell.text)
                text = "\n".join(full_text)
                
        except Exception as e:
            return f"Error processing {filepath}: {e}"
        
        return text.lower()  # Normalize to lowercase

    def extract_experience(self, text):
        """
        Extract total years of experience from resume text
        Returns: float (years of experience) or None if not found
        """
        patterns = [
            r'(\d+\.?\d*)\s*(years?|yrs?)\s*(experience|exp)',
            r'experience\s*:\s*(\d+\.?\d*)\s*(years?|yrs?)',
            r'(\d+)\+?\s*(years?|yrs?)\s*in\s*.*(experience|exp)',
            r'(\d+\.?\d*)\s*(years?|yrs?)\s*professional',
            r'(\d+)\s*(years?|yrs?)\s*relevant'
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                try:
                    return float(match.group(1))
                except:
                    continue
        
        date_pattern = r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s*\d{4}\s*[-–—]\s*(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s*\d{4}|\bpresent\b'
        date_ranges = re.findall(date_pattern, text, flags=re.IGNORECASE)
        
        total_experience = 0
        valid_ranges = 0
        
        for date_range in date_ranges:
            try:
                dates = re.split(r'\s*[-–—]\s*', date_range)
                if len(dates) != 2:
                    continue
                    
                start_date = dates[0].strip()
                end_date = dates[1].strip()
                
                if end_date.lower() == "present":
                    end_date = datetime.now().strftime("%b %Y")
                
                start = datetime.strptime(start_date[:3] + " " + start_date[-4:], "%b %Y")
                end = datetime.strptime(end_date[:3] + " " + end_date[-4:], "%b %Y")
                
                duration = (end - start).days / 365.25
                if duration > 0:
                    total_experience += duration
                    valid_ranges += 1
            except Exception:
                continue
        
        if valid_ranges > 0:
            return total_experience
        
        return None

    def analyze_resume(self, text):
        """
        Analyze entire resume text and return:
        - missing_mandatory: List of missing mandatory keywords
        - keyword_counts: Dictionary of keyword frequencies
        - experience: Extracted years of experience
        - score: Total weighted score
        """
        missing_mandatory = []
        keyword_counts = defaultdict(int)
        found_sections = defaultdict(list)
        
        for kw in self.mandatory:
            pattern = r'\b' + re.escape(kw) + r'\b'
            matches = re.finditer(pattern, text)
            count = 0
            
            for match in matches:
                count += 1
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 50)
                context = text[start:end].replace('\n', ' ').strip()
                found_sections[kw].append(context)
            
            if count == 0:
                missing_mandatory.append(kw)
            keyword_counts[kw] = count
        
        for kw in self.optional:
            pattern = r'\b' + re.escape(kw) + r'\b'
            keyword_counts[kw] = len(re.findall(pattern, text))
        
        experience = self.extract_experience(text)
        
        mandatory_score = sum(keyword_counts[kw] * 3 for kw in self.mandatory)
        optional_score = sum(keyword_counts[kw] for kw in self.optional)
        
        exp_bonus = 0
        if experience is not None and experience >= self.min_experience:
            exp_bonus = min(experience / self.min_experience, 1.2) * (mandatory_score + optional_score) * 0.2
        
        total_score = mandatory_score + optional_score + exp_bonus
        
        return {
            'missing_mandatory': missing_mandatory,
            'keyword_counts': dict(keyword_counts),
            'found_sections': dict(found_sections),
            'experience': experience,
            'score': total_score
        }

    def process_resumes(self, resume_file_paths):
        """Process resumes from file paths and return ranked results"""
        results = []
        
        for filepath in resume_file_paths:
            if not filepath.lower().endswith(('.pdf', '.docx', '.doc')):
                continue
                
            text = self.extract_text(filepath)
            
            if isinstance(text, str) and not text.strip():
                continue
            elif not isinstance(text, str):
                continue
                
            analysis = self.analyze_resume(text)
            
            if not analysis['missing_mandatory']:
                exp_met = analysis['experience'] is not None and analysis['experience'] >= self.min_experience
                
                results.append({
                    'filename': os.path.basename(filepath),
                    'score': analysis['score'],
                    'experience': analysis['experience'],
                    'experience_met': exp_met,
                    'missing_mandatory': analysis['missing_mandatory'],
                    'keyword_counts': analysis['keyword_counts'],
                    'found_sections': analysis['found_sections']
                })
        
        results.sort(key=lambda x: (-x['score'], -x['experience_met']))
        return results

def process_resume_filter(job_description, mandatory_keywords, optional_keywords, min_experience, resume_files):
    """Process inputs from Gradio interface and return formatted results"""
    try:
        # Convert comma-separated strings to lists
        mandatory_keywords = [kw.strip() for kw in mandatory_keywords.split(',') if kw.strip()]
        optional_keywords = [kw.strip() for kw in optional_keywords.split(',') if kw.strip()]
        min_experience = float(min_experience)
        
        # Initialize and run filter
        resume_filter = ResumeFilter(job_description, mandatory_keywords, optional_keywords, min_experience)
        
        # Handle resume_files as file paths
        resume_file_paths = []
        if resume_files:
            if isinstance(resume_files, str):  # Single file case
                resume_file_paths = [resume_files]
            elif isinstance(resume_files, list):  # Multiple files
                resume_file_paths = [f.name if hasattr(f, 'name') else f for f in resume_files]
        
        results = resume_filter.process_resumes(resume_file_paths)
        
        # Format output
        output = f"{'='*50}\nJob Requirements:\n{job_description}\n{'='*50}\n\n"
        output += f"Minimum Experience Required: {min_experience} years\n\n"
        
        if not results:
            output += "No resumes matched all requirements"
            return output
        
        output += f"Found {len(results)} qualified resumes:\n\n"
        for rank, resume in enumerate(results, 1):
            exp_status = "✓" if resume['experience_met'] else f"✗ (Found: {resume['experience'] or 'N/A'} yrs)"
            
            output += f"{rank}. {resume['filename']} (Score: {resume['score']:.1f})\n"
            output += f"   Experience: {exp_status}\n"
            output += f"   Missing Mandatory: {', '.join(resume['missing_mandatory']) or 'None'}\n"
            
            output += "   Keyword Evidence:\n"
            for kw in mandatory_keywords:
                count = resume['keyword_counts'].get(kw, 0)
                if count > 0:
                    output += f"     - {kw} ({count} occurrences):\n"
                    for i, context in enumerate(resume['found_sections'].get(kw, [])[:2], 1):
                        output += f"       {i}. {context[:100]}...\n"
            output += "\n" + "-"*50 + "\n"
        
        return output
    
    except Exception as e:
        return f"Error: {str(e)}"

# Gradio Interface
with gr.Blocks() as demo:
    gr.Markdown("# Resume Filter Application")
    
    with gr.Row():
        with gr.Column():
            job_description = gr.Textbox(
                label="Job Description",
                lines=5,
                value="""

                """
            )
            mandatory_keywords = gr.Textbox(
                label="Mandatory Keywords (comma-separated)",
                value=""
            )
            optional_keywords = gr.Textbox(
                label="Optional Keywords (comma-separated)",
                value=""
            )
            min_experience = gr.Textbox(
                label="Minimum Experience (years)",
                value=""
            )
            resume_files = gr.File(
                label="Upload Resumes (PDF or DOCX)",
                file_count="multiple",
                file_types=[".pdf", ".docx", ".doc"]
            )
            submit_button = gr.Button("Process Resumes")
        
        with gr.Column():
            output = gr.Textbox(
                label="Results",
                lines=20,
                interactive=False
            )
    
    submit_button.click(
        fn=process_resume_filter,
        inputs=[job_description, mandatory_keywords, optional_keywords, min_experience, resume_files],
        outputs=output
    )

if __name__ == "__main__":
    demo.launch(share=True)
'''

### Key Changes
1. **Updated `process_resume_filter`**:
   - Modified to handle `resume_files` as a list of file paths or `NamedString` objects.
   - Checks if `resume_files` is a string (single file) or a list (multiple files).
   - Extracts file paths using `f.name` if the file object has a `name` attribute (Gradio’s behavior for uploaded files).
   - Passes file paths directly to `ResumeFilter.process_resumes`.

2. **Updated `process_resumes`**:
   - Changed the parameter to `resume_file_paths` to clarify it expects file paths.
   - Uses `os.path.basename` to get the filename for display purposes.
   - Processes files directly from their paths, eliminating the need to read file contents manually.

3. **Removed FastAPI Dependency**:
   - The previous FastAPI setup was causing complexity and wasn’t necessary for Gradio’s interface.
   - Reverted to a pure Gradio setup, as Gradio handles the web server and file uploads internally.

4. **Error Handling**:
   - Added checks for empty or invalid `resume_files` to prevent errors.
   - Returns meaningful error messages if processing fails.

### Testing Locally
1. Ensure dependencies are installed:
   ```bash
   pip install gradio pdfplumber pytesseract python-docx pdf2image
   ```
2. Install Tesseract and Poppler for OCR (if needed):
   - On Ubuntu: `sudo apt-get install tesseract-ocr poppler-utils`
   - On macOS: `brew install tesseract poppler`
   - On Windows: Install Tesseract and Poppler binaries and add to PATH.
3. Run the script:
   ```bash
   python resume_filter_gui.py
   ```
4. Open the Gradio interface (usually `http://localhost:7860`), upload some PDF/DOCX resumes, and test the form.

### Deploying to Vercel
To host this updated app on Vercel, follow these steps (adjusting from the previous response):

1. **Prepare Project Structure**:
   - Create `requirements.txt`:
     ```
     gradio>=3.50.2
     pdfplumber>=0.10.2
     pytesseract>=0.3.10
     python-docx>=1.0.0
     pdf2image>=1.16.3

     '''